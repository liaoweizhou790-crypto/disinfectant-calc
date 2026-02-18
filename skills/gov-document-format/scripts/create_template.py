#!/usr/bin/env python3
"""
党政机关公文模板生成器
依据 GB/T 9704-2012 标准
"""

from docx import Document
from docx.shared import Mm, Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def set_run_font(run, font_name, font_size, bold=False, color=None):
    """设置字体"""
    run.font.name = font_name
    run.font.size = Pt(font_size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = color
    # 设置中文字体
    r = run._element
    rPr = r.get_or_add_rPr()
    rFonts = rPr.get_or_add_rFonts()
    rFonts.set(qn('w:eastAsia'), font_name)


def set_paragraph_format(paragraph, line_spacing=None, space_before=0, space_after=0, 
                         first_line_indent=None, alignment=None):
    """设置段落格式"""
    if line_spacing:
        paragraph.paragraph_format.line_spacing = Pt(line_spacing)
    paragraph.paragraph_format.space_before = Pt(space_before)
    paragraph.paragraph_format.space_after = Pt(space_after)
    if first_line_indent:
        paragraph.paragraph_format.first_line_indent = Pt(first_line_indent)
    if alignment:
        paragraph.paragraph_format.alignment = alignment


def create_gov_document(output_path, doc_type="通知"):
    """创建标准公文模板"""
    doc = Document()
    
    # 设置页面
    section = doc.sections[0]
    section.page_width = Mm(210)
    section.page_height = Mm(297)
    section.top_margin = Mm(37)
    section.bottom_margin = Mm(35)
    section.left_margin = Mm(28)
    section.right_margin = Mm(26)
    
    # 设置默认正文字体
    style = doc.styles['Normal']
    style.font.name = '仿宋_GB2312'
    style.font.size = Pt(16)
    style._element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
    
    # 1. 发文机关标志（红色，居中）
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("××市人民政府文件")
    set_run_font(run, '方正小标宋简体', 22, color=RGBColor(0xFF, 0x00, 0x00))
    set_paragraph_format(p, space_after=12)
    
    # 2. 发文字号
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("×政发〔2026〕1号")
    set_run_font(run, '仿宋_GB2312', 16)
    set_paragraph_format(p, space_after=18)
    
    # 3. 红色分隔线（用段落边框模拟）
    p = doc.add_paragraph()
    set_paragraph_format(p, space_after=18)
    
    # 4. 标题
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title = f"关于××××××××的{doc_type}"
    run = p.add_run(title)
    set_run_font(run, '方正小标宋简体', 22)
    set_paragraph_format(p, line_spacing=29, space_after=18)
    
    # 5. 主送机关
    p = doc.add_paragraph()
    run = p.add_run("各区、县人民政府，市政府各委、办、局，各市属机构：")
    set_run_font(run, '仿宋_GB2312', 16)
    set_paragraph_format(p, line_spacing=29, space_after=0)
    
    # 6. 正文
    body_text = """　　正文内容正文内容正文内容正文内容正文内容正文内容正文内容正文内容正文内容正文内容正文内容正文内容正文内容正文内容正文内容正文内容正文内容正文内容。

　　一、一级标题

　　（一）二级标题

　　1.三级标题

　　正文内容正文内容正文内容正文内容正文内容正文内容正文内容正文内容正文内容正文内容正文内容正文内容正文内容正文内容正文内容正文内容。

　　二、工作要求

　　（一）加强组织领导
　　正文内容正文内容正文内容正文内容正文内容正文内容正文内容。

　　（二）强化责任落实
　　正文内容正文内容正文内容正文内容正文内容正文内容正文内容正文内容正文内容正文内容。"""
    
    paragraphs = body_text.split('\n\n')
    for para_text in paragraphs:
        if para_text.strip():
            p = doc.add_paragraph()
            # 检查是否是标题行
            if para_text.strip().startswith('　　一、') or para_text.strip().startswith('　　二、') or \
               para_text.strip().startswith('　　三、') or para_text.strip().startswith('　　四、') or \
               para_text.strip().startswith('　　五、'):
                run = p.add_run(para_text.strip())
                set_run_font(run, '黑体', 16, bold=True)
            elif para_text.strip().startswith('　　（'):
                run = p.add_run(para_text.strip())
                set_run_font(run, '楷体_GB2312', 16, bold=True)
            else:
                run = p.add_run(para_text.strip())
                set_run_font(run, '仿宋_GB2312', 16)
            set_paragraph_format(p, line_spacing=29, first_line_indent=32)
    
    # 7. 附件说明
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Pt(32)
    run = p.add_run("附件：1.××××××××××××")
    set_run_font(run, '仿宋_GB2312', 16)
    set_paragraph_format(p, line_spacing=29, space_before=12, space_after=12)
    
    # 8. 发文机关署名和成文日期（留出盖章位置）
    for _ in range(3):
        doc.add_paragraph()
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run("××市人民政府")
    set_run_font(run, '仿宋_GB2312', 16)
    set_paragraph_format(p, space_after=6)
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run("2026年2月13日")
    set_run_font(run, '仿宋_GB2312', 16)
    
    # 9. 附注（如有）
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Pt(32)
    run = p.add_run("（此件公开发布）")
    set_run_font(run, '仿宋_GB2312', 16)
    set_paragraph_format(p, space_before=12)
    
    # 10. 版记（抄送、印发机关等）
    doc.add_page_break()
    
    # 版记分隔线
    for _ in range(2):
        doc.add_paragraph()
    
    p = doc.add_paragraph()
    run = p.add_run("抄送：市委各部门，市人大常委会办公厅，市政协办公厅，市法院，市检察院。")
    set_run_font(run, '仿宋_GB2312', 14)
    
    p = doc.add_paragraph()
    run = p.add_run("××市人民政府办公厅")
    run2 = p.add_run("　　　　　　　　　　　　　　　　2026年2月13日印发")
    set_run_font(run, '仿宋_GB2312', 14)
    set_run_font(run2, '仿宋_GB2312', 14)
    
    # 保存
    doc.save(output_path)
    print(f"公文模板已生成：{output_path}")


if __name__ == "__main__":
    import sys
    output = sys.argv[1] if len(sys.argv) > 1 else "公文模板.docx"
    doc_type = sys.argv[2] if len(sys.argv) > 2 else "通知"
    create_gov_document(output, doc_type)
